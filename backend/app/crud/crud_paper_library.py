"""
试卷库功能相关的CRUD操作
"""

from sqlalchemy.orm import Session, joinedload
from sqlalchemy import or_, and_, func
from typing import Optional, List, Dict, Any
from datetime import datetime, timezone, timedelta
import tempfile
import os
import json

from app.models import models
from app.schemas import schemas


def get_paper_library_list(
    db: Session,
    teacher_id: int,
    keyword: Optional[str] = None,
    direction: Optional[str] = None,
    difficulty: Optional[str] = None,
    source: Optional[str] = None,
    skip: int = 0,
    limit: int = 20
):
    """获取试卷库列表"""
    # 基础查询
    query = db.query(models.TestPaper).filter(
        or_(
            models.TestPaper.creator_id == teacher_id,  # 个人试卷
            models.TestPaper.is_shared == True  # 共享试卷
        ),
        models.TestPaper.deleted_at.is_(None)  # 使用deleted_at字段
    )
    
    # 关键词搜索
    if keyword:
        query = query.filter(models.TestPaper.title.ilike(f"%{keyword}%"))
    
    # 方向筛选
    if direction:
        query = query.filter(models.TestPaper.direction == direction)
    
    # 难度筛选
    if difficulty:
        query = query.filter(models.TestPaper.difficulty == difficulty)
    
    # 来源筛选
    if source:
        if source == "personal":
            query = query.filter(models.TestPaper.creator_id == teacher_id)
        elif source == "platform":
            query = query.filter(models.TestPaper.is_shared == True)
    
    # 获取总数
    total = query.count()
    
    # 分页查询
    papers = query.order_by(
        models.TestPaper.created_at.desc()
    ).offset(skip).limit(limit).all()
    
    # 构建响应数据
    paper_list = []
    for paper in papers:
        # 获取试题数量
        question_count = db.query(models.TestPaperQuestion).filter(
            models.TestPaperQuestion.test_paper_id == paper.id
        ).count()
        
        # 判断来源
        source_type = "个人" if paper.creator_id == teacher_id else "平台"
        
        # 难度中文显示
        difficulty_cn_map = {
            "BEGINNER": "初级",
            "INTERMEDIATE": "中级", 
            "ADVANCED": "高级"
        }
        difficulty_cn = difficulty_cn_map.get(paper.difficulty.value if paper.difficulty else None, "未设置")
        
        # 获取创建者信息
        creator = None
        if paper.creator_id:
            creator = db.query(models.User).filter(models.User.id == paper.creator_id).first()
        
        paper_data = {
            "id": paper.id,
            "title": paper.title,
            "direction": paper.direction,
            "difficulty": paper.difficulty.value if paper.difficulty else None,
            "difficulty_cn": difficulty_cn,
            "question_count": question_count,
            "total_score": paper.total_score,
            "source": source_type,
            "creator_name": creator.full_name if creator else "系统",
            "created_at": paper.created_at,
            "can_create_exam": True,
            "can_copy": True,
            "can_export": True,
            "can_edit": paper.creator_id == teacher_id,
            "can_delete": paper.creator_id == teacher_id
        }
        paper_list.append(paper_data)
    
    return paper_list, total


def create_test_paper(
    db: Session,
    paper_data: dict,
    creator_id: int
):
    """创建试卷"""
    # 创建试卷记录
    new_paper = models.TestPaper(
        title=paper_data["title"],
        description=paper_data.get("description"),
        direction=paper_data.get("direction"),
        difficulty=paper_data.get("difficulty"),
        composition_method=paper_data.get("composition_method", "MANUAL_SELECTION"),
        creator_id=creator_id,
        is_shared=False,
        total_score=0  # 初始分数为0，后续添加题目时更新
    )
    
    db.add(new_paper)
    db.commit()
    db.refresh(new_paper)
    
    return new_paper


def copy_test_paper(
    db: Session,
    paper_id: int,
    teacher_id: int,
    new_title: Optional[str] = None
):
    """复制试卷"""
    # 获取原试卷
    original_paper = db.query(models.TestPaper).filter(
        models.TestPaper.id == paper_id,
        models.TestPaper.deleted_at.is_(None)  # 使用deleted_at字段
    ).first()
    
    if not original_paper:
        return None
    
    # 生成新标题
    if not new_title:
        new_title = f"{original_paper.title}_副本"
        # 如果超过60字符，截断但仍允许保存
        if len(new_title) > 60:
            new_title = new_title[:57] + "..."
    
    # 创建新试卷
    new_paper = models.TestPaper(
        title=new_title,
        description=original_paper.description,
        direction=original_paper.direction,
        difficulty=original_paper.difficulty,
        composition_method=original_paper.composition_method,
        creator_id=teacher_id,
        is_shared=False,
        total_score=original_paper.total_score
    )
    
    db.add(new_paper)
    db.flush()  # 获取新试卷ID
    
    # 复制试题关联
    original_questions = db.query(models.TestPaperQuestion).filter(
        models.TestPaperQuestion.test_paper_id == paper_id
    ).all()
    
    for orig_q in original_questions:
        new_question = models.TestPaperQuestion(
            test_paper_id=new_paper.id,
            question_id=orig_q.question_id,
            order_in_paper=orig_q.order_in_paper,
            score_for_question=orig_q.score_for_question
        )
        db.add(new_question)
    
    db.commit()
    db.refresh(new_paper)
    
    return new_paper


def delete_test_paper(
    db: Session,
    paper_id: int,
    teacher_id: int
):
    """删除试卷（仅个人试卷）"""
    # 获取试卷
    paper = db.query(models.TestPaper).filter(
        models.TestPaper.id == paper_id,
        models.TestPaper.creator_id == teacher_id,  # 只能删除自己创建的试卷
        models.TestPaper.deleted_at.is_(None)  # 使用deleted_at字段
    ).first()
    
    if not paper:
        return False
    
    # 软删除
    paper.deleted_at = datetime.now(timezone.utc)
    db.commit()
    
    return True


def update_test_paper(
    db: Session,
    paper_id: int,
    teacher_id: int,
    update_data: dict
):
    """更新试卷信息"""
    # 获取试卷
    paper = db.query(models.TestPaper).filter(
        models.TestPaper.id == paper_id,
        models.TestPaper.creator_id == teacher_id,  # 只能编辑自己创建的试卷
        models.TestPaper.deleted_at.is_(None)  # 使用deleted_at字段
    ).first()
    
    if not paper:
        return None
    
    # 更新字段
    for field, value in update_data.items():
        if hasattr(paper, field):
            setattr(paper, field, value)
    
    db.commit()
    db.refresh(paper)
    
    return paper


def get_test_paper_detail(
    db: Session,
    paper_id: int,
    teacher_id: int
):
    """获取试卷详情"""
    # 获取试卷基本信息
    paper = db.query(models.TestPaper).filter(
        models.TestPaper.id == paper_id,
        models.TestPaper.deleted_at.is_(None),  # 使用deleted_at字段
        or_(
            models.TestPaper.creator_id == teacher_id,
            models.TestPaper.is_shared == True
        )
    ).first()
    
    if not paper:
        return None
    
    # 获取创建者信息
    creator = None
    if paper.creator_id:
        creator = db.query(models.User).filter(models.User.id == paper.creator_id).first()
    
    # 获取试题列表
    questions = db.query(models.TestPaperQuestion).options(
        joinedload(models.TestPaperQuestion.question)
    ).filter(
        models.TestPaperQuestion.test_paper_id == paper_id
    ).order_by(models.TestPaperQuestion.order_in_paper).all()
    
    # 构建试题数据
    question_list = []
    for tpq in questions:
        question = tpq.question

        # 解析options JSON字符串
        options = []
        if question.options:
            try:
                import json
                options = json.loads(question.options)
            except (json.JSONDecodeError, TypeError):
                options = []

        # 解析correct_answers JSON字符串
        correct_answers = []
        if question.correct_answers:
            try:
                import json
                correct_answers = json.loads(question.correct_answers)
            except (json.JSONDecodeError, TypeError):
                correct_answers = []

        question_data = {
            "id": question.id,
            "type": question.question_type.value if question.question_type else "SINGLE_CHOICE",
            "content": question.content,
            "score": tpq.score_for_question,
            "order_index": tpq.order_in_paper,
            "options": options,
            "correct_answers": correct_answers,
            "explanation": question.explanation
        }
        question_list.append(question_data)
    
    # 难度中文显示
    difficulty_cn_map = {
        "BEGINNER": "初级",
        "INTERMEDIATE": "中级", 
        "ADVANCED": "高级"
    }
    difficulty_cn = difficulty_cn_map.get(paper.difficulty.value if paper.difficulty else None, "未设置")
    
    # 构建试卷详情
    paper_detail = {
        "id": paper.id,
        "title": paper.title,
        "description": paper.description,
        "direction": paper.direction,
        "difficulty": paper.difficulty.value if paper.difficulty else None,
        "difficulty_cn": difficulty_cn,
        "composition_method": paper.composition_method,
        "total_score": paper.total_score,
        "estimated_duration_minutes": paper.estimated_duration_minutes,
        "question_count": len(question_list),
        "is_shared": paper.is_shared,
        "creator_id": paper.creator_id,
        "creator_name": creator.full_name if creator else "系统",
        "created_at": paper.created_at,
        "updated_at": paper.updated_at,
        "can_edit": paper.creator_id == teacher_id,
        "can_delete": paper.creator_id == teacher_id,
        "questions": question_list
    }
    
    return paper_detail


def create_exam_from_paper(
    db: Session,
    paper_id: int,
    classroom_id: int,
    teacher_id: int,
    exam_title: str
):
    """从试卷创建考试"""
    # 将导入移到函数内部避免循环导入
    from app.crud.crud import check_classroom_teacher_permission

    # 先检查课堂是否存在（避免 FK 违反导致 500；超级管理员权限检查会绕过此校验）
    classroom = db.query(models.Classroom).filter(
        models.Classroom.id == classroom_id
    ).first()
    if not classroom:
        return None

    # 检查权限
    if not check_classroom_teacher_permission(db, classroom_id, teacher_id):
        return None
    
    # 检查试卷是否存在
    paper = db.query(models.TestPaper).filter(
        models.TestPaper.id == paper_id,
        models.TestPaper.deleted_at.is_(None),  # 使用deleted_at字段
        or_(
            models.TestPaper.creator_id == teacher_id,
            models.TestPaper.is_shared == True
        )
    ).first()
    
    if not paper:
        return None
    
    # 创建考试（默认 end = start + 2 小时，确保默认值可发布）
    _now = datetime.now(timezone.utc)
    new_exam = models.ClassroomExam(
        classroom_id=classroom_id,
        test_paper_id=paper_id,
        title=exam_title,
        created_by_teacher_id=teacher_id,
        status=models.ExamStatusEnum.UNPUBLISHED,
        exam_start_time=_now,  # 默认现在，发布时可修改
        exam_end_time=_now + timedelta(hours=2),  # 默认 2 小时后，确保 end > start
        duration_minutes=60  # 默认60分钟，后续发布时修改
    )
    
    db.add(new_exam)
    db.commit()
    db.refresh(new_exam)
    
    return new_exam


def get_paper_filter_tags(db: Session):
    """获取试卷筛选标签"""
    # 获取所有方向
    directions = db.query(models.TestPaper.direction).filter(
        models.TestPaper.direction.isnot(None),
        models.TestPaper.deleted_at.is_(None)  # 使用deleted_at字段
    ).distinct().all()
    
    # 获取所有难度
    difficulties = db.query(models.TestPaper.difficulty).filter(
        models.TestPaper.difficulty.isnot(None),
        models.TestPaper.deleted_at.is_(None)  # 使用deleted_at字段
    ).distinct().all()
    
    # 难度标签映射
    difficulty_map = {
        "BEGINNER": {"value": "BEGINNER", "label": "初级"},
        "INTERMEDIATE": {"value": "INTERMEDIATE", "label": "中级"},
        "ADVANCED": {"value": "ADVANCED", "label": "高级"}
    }
    
    difficulty_tags = []
    for d in difficulties:
        if d[0] and d[0].value in difficulty_map:
            difficulty_tags.append(difficulty_map[d[0].value])
    
    return {
        "directions": [d[0] for d in directions if d[0]],
        "difficulties": difficulty_tags,
        "sources": [
            {"value": "personal", "label": "个人"},
            {"value": "platform", "label": "平台"}
        ]
    }


def export_test_paper_to_word(
    db: Session,
    paper_id: int,
    teacher_id: int
):
    """导出试卷为Word文档"""
    # 获取试卷详情
    paper_detail = get_test_paper_detail(db, paper_id, teacher_id)
    if not paper_detail:
        return None
    
    try:
        from docx import Document
        from docx.shared import Inches
        
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        title = doc.add_heading(paper_detail["title"], 0)
        title.alignment = 1  # 居中对齐
        
        # 添加试卷信息
        info_para = doc.add_paragraph()
        info_para.add_run(f"方向：{paper_detail['direction'] or '未设置'}")
        info_para.add_run(f"    难度：{paper_detail['difficulty_cn'] or '未设置'}")
        info_para.add_run(f"    总分：{paper_detail['total_score']}分")
        info_para.add_run(f"    题目数量：{paper_detail['question_count']}题")
        
        # 添加分隔线
        doc.add_paragraph("=" * 50)
        
        # 添加题目
        for i, question in enumerate(paper_detail["questions"], 1):
            # 题目标题
            q_para = doc.add_paragraph()
            q_para.add_run(f"{i}. ").bold = True
            q_para.add_run(f"({question['score']}分) ")
            q_para.add_run(question["content"])
            
            # 选择题选项
            if question["type"] in ["SINGLE_CHOICE", "MULTIPLE_CHOICE"] and question["options"]:
                try:
                    import json
                    options = json.loads(question["options"]) if isinstance(question["options"], str) else question["options"]
                    for j, option in enumerate(options):
                        option_para = doc.add_paragraph()
                        option_text = option if isinstance(option, str) else option.get("text", "")
                        option_para.add_run(f"    {chr(65+j)}. {option_text}")
                except:
                    pass
            
            # 添加空行
            doc.add_paragraph()
        
        # 保存到临时文件
        temp_dir = tempfile.gettempdir()
        filename = f"试卷_{paper_detail['title']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(temp_dir, filename)
        
        doc.save(filepath)
        
        return {
            "filename": filename,
            "filepath": filepath,
            "success": True
        }
        
    except ImportError:
        return {
            "success": False,
            "error": "缺少python-docx依赖，无法导出Word文档"
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"导出失败：{str(e)}"
        }


def get_user_classrooms_for_exam(
    db: Session,
    teacher_id: int
):
    """获取教师的课堂列表，用于创建考试时选择"""
    classrooms = db.query(models.Classroom).filter(
        models.Classroom.teacher_id == teacher_id,
        models.Classroom.deleted_at.is_(None)  # 使用deleted_at字段
    ).order_by(models.Classroom.created_at.desc()).all()
    
    classroom_list = []
    for classroom in classrooms:
        classroom_data = {
            "id": classroom.id,
            "name": classroom.name,
            "status": classroom.status.value,
            "student_count": classroom.student_count
        }
        classroom_list.append(classroom_data)
    
    return classroom_list

# ==================== 试卷模板相关CRUD操作 ====================

def get_paper_templates(
    db: Session,
    teacher_id: int
):
    """获取试卷模板列表"""
    # 获取系统内置模板
    system_templates = db.query(models.PaperTemplate).options(
        joinedload(models.PaperTemplate.question_configs)
    ).filter(
        models.PaperTemplate.is_system_template == True,
        models.PaperTemplate.is_active == True
    ).order_by(models.PaperTemplate.created_at.asc()).all()
    
    # 获取个人自定义模板
    custom_templates = db.query(models.PaperTemplate).options(
        joinedload(models.PaperTemplate.question_configs),
        joinedload(models.PaperTemplate.creator)
    ).filter(
        models.PaperTemplate.creator_id == teacher_id,
        models.PaperTemplate.is_active == True
    ).order_by(models.PaperTemplate.created_at.desc()).all()
    
    def format_template(template):
        """格式化模板数据"""
        total_questions = sum(config.question_count for config in template.question_configs)
        total_score = sum(config.question_count * config.score_per_question for config in template.question_configs)
        
        question_configs = []
        for config in sorted(template.question_configs, key=lambda x: x.order_index):
            question_configs.append({
                "question_type": config.question_type.value,
                "question_count": config.question_count,
                "score_per_question": config.score_per_question,
                "difficulty": config.difficulty.value if config.difficulty else None,
                "order_index": config.order_index
            })
        
        return {
            "id": template.id,
            "name": template.name,
            "description": template.description,
            "is_system_template": template.is_system_template,
            "creator_id": template.creator_id,
            "creator_name": template.creator.full_name if template.creator else "系统",
            "question_configs": question_configs,
            "total_questions": total_questions,
            "total_score": total_score,
            "created_at": template.created_at,
            "updated_at": template.updated_at
        }
    
    return {
        "system_templates": [format_template(t) for t in system_templates],
        "custom_templates": [format_template(t) for t in custom_templates]
    }


def create_paper_template(
    db: Session,
    template_data: dict,
    creator_id: int
):
    """创建试卷模板"""
    # 创建模板记录
    new_template = models.PaperTemplate(
        name=template_data["name"],
        description=template_data.get("description"),
        creator_id=creator_id,
        is_system_template=False,
        is_active=True
    )
    
    db.add(new_template)
    db.flush()  # 获取模板ID
    
    # 创建题目配置
    for config_data in template_data["question_configs"]:
        config = models.PaperTemplateQuestionConfig(
            template_id=new_template.id,
            question_type=config_data["question_type"],
            question_count=config_data["question_count"],
            score_per_question=config_data["score_per_question"],
            difficulty=config_data.get("difficulty"),
            order_index=config_data.get("order_index", 0)
        )
        db.add(config)
    
    db.commit()
    db.refresh(new_template)
    
    return new_template


def get_question_pool_stats(
    db: Session,
    teacher_id: int,
    direction: Optional[str] = None,
    source: Optional[str] = None
):
    """获取题库统计信息"""
    # 基础查询
    query = db.query(models.Question).filter(
        or_(
            models.Question.creator_id == teacher_id,  # 个人题目
            models.Question.is_shared == True  # 共享题目
        )
    )
    
    # 方向筛选
    if direction:
        query = query.filter(models.Question.direction == direction)
    
    # 来源筛选
    if source:
        if source == "personal":
            query = query.filter(models.Question.creator_id == teacher_id)
        elif source == "platform":
            query = query.filter(models.Question.is_shared == True)
    
    questions = query.all()
    
    # 统计各题型数量
    question_type_stats = {}
    difficulty_stats = {}
    
    for question in questions:
        # 题型统计
        q_type = question.question_type.value
        question_type_stats[q_type] = question_type_stats.get(q_type, 0) + 1
        
        # 难度统计
        difficulty = question.difficulty.value if question.difficulty else "INTERMEDIATE"
        difficulty_stats[difficulty] = difficulty_stats.get(difficulty, 0) + 1
    
    return {
        "total_questions": len(questions),
        "question_type_stats": question_type_stats,
        "difficulty_stats": difficulty_stats
    }


def generate_paper_from_template(
    db: Session,
    paper_data: dict,
    template_id: int,
    teacher_id: int,
    direction: Optional[str] = None,
    source: Optional[str] = None
):
    """根据模板生成试卷"""
    import random
    
    # 获取模板
    template = db.query(models.PaperTemplate).options(
        joinedload(models.PaperTemplate.question_configs)
    ).filter(
        models.PaperTemplate.id == template_id,
        models.PaperTemplate.is_active == True
    ).first()
    
    if not template:
        raise ValueError("模板不存在")
    
    # 创建试卷
    new_paper = models.TestPaper(
        title=paper_data["title"],
        description=paper_data.get("description"),
        direction=direction,
        difficulty=paper_data.get("difficulty"),
        composition_method="TEMPLATE_BASED",
        creator_id=teacher_id,
        is_shared=False,
        total_score=0  # 后续计算
    )
    
    db.add(new_paper)
    db.flush()  # 获取试卷ID
    
    total_score = 0
    order_index = 1
    
    # 根据模板配置抽取题目
    for config in sorted(template.question_configs, key=lambda x: x.order_index):
        # 构建题目查询
        query = db.query(models.Question).filter(
            models.Question.question_type == config.question_type,
            or_(
                models.Question.creator_id == teacher_id,
                models.Question.is_shared == True
            )
        )
        
        # 方向筛选
        if direction:
            query = query.filter(models.Question.direction == direction)
        
        # 来源筛选
        if source:
            if source == "personal":
                query = query.filter(models.Question.creator_id == teacher_id)
            elif source == "platform":
                query = query.filter(models.Question.is_shared == True)
        
        # 难度筛选
        if config.difficulty:
            query = query.filter(models.Question.difficulty == config.difficulty)
        
        # 获取可用题目
        available_questions = query.all()
        
        if len(available_questions) < config.question_count:
            raise ValueError(f"题库中{config.question_type.value}类型题目数量不足，需要{config.question_count}道，实际只有{len(available_questions)}道")
        
        # 随机抽取题目
        selected_questions = random.sample(available_questions, config.question_count)
        
        # 添加到试卷
        for question in selected_questions:
            paper_question = models.TestPaperQuestion(
                test_paper_id=new_paper.id,
                question_id=question.id,
                score_for_question=config.score_per_question,
                order_in_paper=order_index
            )
            db.add(paper_question)
            total_score += config.score_per_question
            order_index += 1
    
    # 更新试卷总分
    new_paper.total_score = total_score
    
    db.commit()
    db.refresh(new_paper)
    
    return new_paper


def add_questions_to_paper(
    db: Session,
    paper_id: int,
    question_ids: List[int],
    teacher_id: int,
    scores: Optional[List[int]] = None
):
    """向试卷添加题目"""
    # 验证试卷权限
    paper = db.query(models.TestPaper).filter(
        models.TestPaper.id == paper_id,
        models.TestPaper.creator_id == teacher_id,
        models.TestPaper.deleted_at.is_(None)
    ).first()
    
    if not paper:
        raise ValueError("试卷不存在或无权限编辑")
    
    # 获取当前试卷最大排序号
    max_order = db.query(func.max(models.TestPaperQuestion.order_in_paper)).filter(
        models.TestPaperQuestion.test_paper_id == paper_id
    ).scalar() or 0
    
    # 验证题目存在性和权限
    questions = db.query(models.Question).filter(
        models.Question.id.in_(question_ids),
        or_(
            models.Question.creator_id == teacher_id,
            models.Question.is_shared == True
        )
    ).all()
    
    if len(questions) != len(question_ids):
        raise ValueError("部分题目不存在或无权限访问")
    
    # 检查题目是否已在试卷中
    existing_questions = db.query(models.TestPaperQuestion.question_id).filter(
        models.TestPaperQuestion.test_paper_id == paper_id,
        models.TestPaperQuestion.question_id.in_(question_ids)
    ).all()
    
    if existing_questions:
        existing_ids = [q[0] for q in existing_questions]
        raise ValueError(f"题目ID {existing_ids} 已在试卷中")
    
    # 添加题目到试卷
    added_score = 0
    for i, question in enumerate(questions):
        score = scores[i] if scores else 10  # 默认10分
        
        paper_question = models.TestPaperQuestion(
            test_paper_id=paper_id,
            question_id=question.id,
            score_for_question=score,
            order_in_paper=max_order + i + 1
        )
        db.add(paper_question)
        added_score += score
    
    # 更新试卷总分
    paper.total_score += added_score
    
    db.commit()
    
    return len(questions)


def remove_questions_from_paper(
    db: Session,
    paper_id: int,
    question_ids: List[int],
    teacher_id: int
):
    """从试卷移除题目"""
    # 验证试卷权限
    paper = db.query(models.TestPaper).filter(
        models.TestPaper.id == paper_id,
        models.TestPaper.creator_id == teacher_id,
        models.TestPaper.deleted_at.is_(None)
    ).first()
    
    if not paper:
        raise ValueError("试卷不存在或无权限编辑")
    
    # 获取要删除的题目
    paper_questions = db.query(models.TestPaperQuestion).filter(
        models.TestPaperQuestion.test_paper_id == paper_id,
        models.TestPaperQuestion.question_id.in_(question_ids)
    ).all()
    
    if not paper_questions:
        raise ValueError("指定题目不在试卷中")
    
    # 计算减少的分数
    removed_score = sum(pq.score_for_question for pq in paper_questions)
    
    # 删除题目
    for pq in paper_questions:
        db.delete(pq)
    
    # 更新试卷总分
    paper.total_score = max(0, paper.total_score - removed_score)
    
    db.commit()
    
    return len(paper_questions)


def update_paper_question(
    db: Session,
    paper_id: int,
    question_id: int,
    teacher_id: int,
    update_data: dict
):
    """更新试卷中的题目"""
    # 验证试卷权限
    paper = db.query(models.TestPaper).filter(
        models.TestPaper.id == paper_id,
        models.TestPaper.creator_id == teacher_id,
        models.TestPaper.deleted_at.is_(None)
    ).first()
    
    if not paper:
        raise ValueError("试卷不存在或无权限编辑")
    
    # 获取试卷题目
    paper_question = db.query(models.TestPaperQuestion).filter(
        models.TestPaperQuestion.test_paper_id == paper_id,
        models.TestPaperQuestion.question_id == question_id
    ).first()
    
    if not paper_question:
        raise ValueError("题目不在试卷中")
    
    # 更新分数
    if "score" in update_data and update_data["score"] is not None:
        old_score = paper_question.score_for_question
        new_score = update_data["score"]
        paper_question.score_for_question = new_score
        
        # 更新试卷总分
        paper.total_score = paper.total_score - old_score + new_score
    
    # 更新排序
    if "order_index" in update_data and update_data["order_index"] is not None:
        paper_question.order_in_paper = update_data["order_index"]
    
    db.commit()
    
    return paper_question

# ==================== 新增的试卷分值和排序功能 ====================

def update_paper_question_score(
    db: Session,
    paper_id: int,
    question_id: int,
    teacher_id: int,
    new_score: int
):
    """单独更新题目分值"""
    # 验证试卷权限
    paper = db.query(models.TestPaper).filter(
        models.TestPaper.id == paper_id,
        models.TestPaper.creator_id == teacher_id,
        models.TestPaper.deleted_at.is_(None)
    ).first()
    
    if not paper:
        raise ValueError("试卷不存在或无权限编辑")
    
    # 获取试卷题目
    paper_question = db.query(models.TestPaperQuestion).filter(
        models.TestPaperQuestion.test_paper_id == paper_id,
        models.TestPaperQuestion.question_id == question_id
    ).first()
    
    if not paper_question:
        raise ValueError("题目不在试卷中")
    
    # 更新分数
    old_score = paper_question.score_for_question
    paper_question.score_for_question = new_score
    
    # 更新试卷总分
    paper.total_score = paper.total_score - old_score + new_score
    
    db.commit()
    
    return {
        "question_id": question_id,
        "old_score": old_score,
        "new_score": new_score,
        "paper_total_score": paper.total_score
    }


def batch_update_paper_questions_score(
    db: Session,
    paper_id: int,
    teacher_id: int,
    question_type: str,
    score_per_question: int
):
    """批量设置同类型题目分值"""
    # 验证试卷权限
    paper = db.query(models.TestPaper).filter(
        models.TestPaper.id == paper_id,
        models.TestPaper.creator_id == teacher_id,
        models.TestPaper.deleted_at.is_(None)
    ).first()
    
    if not paper:
        raise ValueError("试卷不存在或无权限编辑")
    
    # 获取指定类型的题目
    paper_questions = db.query(models.TestPaperQuestion).join(
        models.Question
    ).filter(
        models.TestPaperQuestion.test_paper_id == paper_id,
        models.Question.question_type == question_type
    ).all()
    
    if not paper_questions:
        raise ValueError(f"试卷中没有{question_type}类型的题目")
    
    # 计算分数变化
    total_old_score = sum(pq.score_for_question for pq in paper_questions)
    total_new_score = len(paper_questions) * score_per_question
    
    # 批量更新分数
    updated_count = 0
    for paper_question in paper_questions:
        paper_question.score_for_question = score_per_question
        updated_count += 1
    
    # 更新试卷总分
    paper.total_score = paper.total_score - total_old_score + total_new_score
    
    db.commit()
    
    return {
        "question_type": question_type,
        "updated_count": updated_count,
        "score_per_question": score_per_question,
        "total_score_change": total_new_score - total_old_score,
        "paper_total_score": paper.total_score
    }


def update_paper_questions_order(
    db: Session,
    paper_id: int,
    teacher_id: int,
    question_orders: List[dict]
):
    """更新题目排序（小题排序）"""
    # 验证试卷权限
    paper = db.query(models.TestPaper).filter(
        models.TestPaper.id == paper_id,
        models.TestPaper.creator_id == teacher_id,
        models.TestPaper.deleted_at.is_(None)
    ).first()
    
    if not paper:
        raise ValueError("试卷不存在或无权限编辑")
    
    # 获取要更新的题目
    question_ids = [item['question_id'] for item in question_orders]
    paper_questions = db.query(models.TestPaperQuestion).filter(
        models.TestPaperQuestion.test_paper_id == paper_id,
        models.TestPaperQuestion.question_id.in_(question_ids)
    ).all()
    
    if len(paper_questions) != len(question_ids):
        raise ValueError("部分题目不在试卷中")
    
    # 创建题目ID到排序的映射
    order_map = {item['question_id']: item['order_index'] for item in question_orders}
    
    # 更新排序
    updated_count = 0
    for paper_question in paper_questions:
        new_order = order_map[paper_question.question_id]
        paper_question.order_in_paper = new_order
        updated_count += 1
    
    db.commit()
    
    return {
        "updated_count": updated_count,
        "question_orders": question_orders
    }


def update_paper_sections_order(
    db: Session,
    paper_id: int,
    teacher_id: int,
    section_orders: List[dict]
):
    """更新大题排序"""
    # 验证试卷权限
    paper = db.query(models.TestPaper).filter(
        models.TestPaper.id == paper_id,
        models.TestPaper.creator_id == teacher_id,
        models.TestPaper.deleted_at.is_(None)
    ).first()
    
    if not paper:
        raise ValueError("试卷不存在或无权限编辑")
    
    # 获取试卷中的所有题目，按题型分组
    paper_questions = db.query(models.TestPaperQuestion).join(
        models.Question
    ).filter(
        models.TestPaperQuestion.test_paper_id == paper_id
    ).order_by(models.TestPaperQuestion.order_in_paper).all()
    
    if not paper_questions:
        raise ValueError("试卷中没有题目")
    
    # 按题型分组
    questions_by_type = {}
    for pq in paper_questions:
        question_type = pq.question.question_type.value
        if question_type not in questions_by_type:
            questions_by_type[question_type] = []
        questions_by_type[question_type].append(pq)
    
    # 创建题型到新排序的映射
    type_order_map = {item['question_type']: item['order_index'] for item in section_orders}
    
    # 重新分配排序号
    current_order = 1
    updated_count = 0
    
    # 按新的大题顺序重新排序
    for section_order in sorted(section_orders, key=lambda x: x['order_index']):
        question_type = section_order['question_type']
        if question_type in questions_by_type:
            # 为该题型的所有题目重新分配连续的排序号
            for pq in questions_by_type[question_type]:
                pq.order_in_paper = current_order
                current_order += 1
                updated_count += 1
    
    db.commit()
    
    return {
        "updated_count": updated_count,
        "section_orders": section_orders
    }


def get_paper_enhanced_detail(
    db: Session,
    paper_id: int,
    teacher_id: int
):
    """获取增强的试卷详情（包含分组和统计信息）"""
    # 获取试卷基本信息
    paper = db.query(models.TestPaper).filter(
        models.TestPaper.id == paper_id,
        models.TestPaper.deleted_at.is_(None)
    ).first()
    
    if not paper:
        raise ValueError("试卷不存在")
    
    # 获取试卷题目详情
    paper_questions = db.query(models.TestPaperQuestion).join(
        models.Question
    ).filter(
        models.TestPaperQuestion.test_paper_id == paper_id
    ).order_by(models.TestPaperQuestion.order_in_paper).all()
    
    # 构建题目详情列表
    questions = []
    question_type_stats = {}
    sections = {}
    
    for pq in paper_questions:
        question = pq.question
        question_type = question.question_type.value
        
        # 解析选项和答案
        options = None
        correct_answers = None
        if question.options:
            try:
                import json
                options = json.loads(question.options)
            except:
                options = None
        
        if question.correct_answers:
            try:
                import json
                correct_answers = json.loads(question.correct_answers)
            except:
                correct_answers = None
        
        # 题型中文名映射
        type_cn_map = {
            "SINGLE_CHOICE": "单选题",
            "MULTIPLE_CHOICE": "多选题", 
            "TRUE_FALSE": "判断题",
            "SHORT_ANSWER": "简答题"
        }
        
        # 难度中文名映射
        difficulty_cn_map = {
            "BEGINNER": "初级",
            "INTERMEDIATE": "中级",
            "ADVANCED": "高级"
        }
        
        question_detail = {
            "id": pq.id,
            "question_id": question.id,
            "content": question.content,
            "question_type": question_type,
            "question_type_cn": type_cn_map.get(question_type, question_type),
            "options": options,
            "correct_answers": correct_answers,
            "explanation": question.explanation,
            "difficulty": question.difficulty.value,
            "difficulty_cn": difficulty_cn_map.get(question.difficulty.value, question.difficulty.value),
            "score_for_question": pq.score_for_question,
            "order_in_paper": pq.order_in_paper,
            "section_title": pq.section_title
        }
        
        questions.append(question_detail)
        
        # 统计信息
        if question_type not in question_type_stats:
            question_type_stats[question_type] = {
                "count": 0,
                "total_score": 0,
                "type_cn": type_cn_map.get(question_type, question_type)
            }
        
        question_type_stats[question_type]["count"] += 1
        question_type_stats[question_type]["total_score"] += pq.score_for_question
        
        # 大题分组信息
        if question_type not in sections:
            sections[question_type] = {
                "question_type": question_type,
                "question_type_cn": type_cn_map.get(question_type, question_type),
                "question_count": 0,
                "total_score": 0,
                "score_per_question": pq.score_for_question,  # 假设同类型题目分值相同
                "questions": []
            }
        
        sections[question_type]["question_count"] += 1
        sections[question_type]["total_score"] += pq.score_for_question
        sections[question_type]["questions"].append(question_detail)
    
    # 构建统计信息
    statistics = {
        "total_score": paper.total_score,
        "question_count": len(questions),
        "question_type_stats": question_type_stats,
        "section_stats": list(sections.values())
    }
    
    # 权限检查
    can_edit = paper.creator_id == teacher_id
    can_delete = paper.creator_id == teacher_id
    
    # 难度中文名
    difficulty_cn_map = {
        "BEGINNER": "初级",
        "INTERMEDIATE": "中级", 
        "ADVANCED": "高级"
    }
    
    # 获取创建者信息
    creator_name = None
    if paper.creator_id:
        creator = db.query(models.User).filter(models.User.id == paper.creator_id).first()
        if creator:
            creator_name = creator.full_name or creator.username
    
    return {
        "id": paper.id,
        "title": paper.title,
        "description": paper.description,
        "direction": paper.direction,
        "difficulty": paper.difficulty.value if paper.difficulty else None,
        "difficulty_cn": difficulty_cn_map.get(paper.difficulty.value if paper.difficulty else None),
        "composition_method": paper.composition_method,
        "total_score": paper.total_score,
        "estimated_duration_minutes": paper.estimated_duration_minutes,
        "question_count": len(questions),
        "is_shared": paper.is_shared,
        "creator_id": paper.creator_id,
        "creator_name": creator_name,
        "created_at": paper.created_at,
        "updated_at": paper.updated_at,
        "sections": list(sections.values()),
        "questions": questions,
        "statistics": statistics,
        "can_edit": can_edit,
        "can_delete": can_delete,
        "can_copy": True,
        "can_export": True
    }


def update_paper_basic_info(
    db: Session,
    paper_id: int,
    teacher_id: int,
    update_data: dict
):
    """更新试卷基本信息"""
    # 验证试卷权限
    paper = db.query(models.TestPaper).filter(
        models.TestPaper.id == paper_id,
        models.TestPaper.creator_id == teacher_id,
        models.TestPaper.deleted_at.is_(None)
    ).first()
    
    if not paper:
        raise ValueError("试卷不存在或无权限编辑")
    
    # 更新字段
    updated_fields = []
    if "title" in update_data and update_data["title"] is not None:
        paper.title = update_data["title"]
        updated_fields.append("title")
    
    if "direction" in update_data and update_data["direction"] is not None:
        paper.direction = update_data["direction"]
        updated_fields.append("direction")
    
    if "difficulty" in update_data and update_data["difficulty"] is not None:
        paper.difficulty = update_data["difficulty"]
        updated_fields.append("difficulty")
    
    if "description" in update_data and update_data["description"] is not None:
        paper.description = update_data["description"]
        updated_fields.append("description")
    
    if "estimated_duration_minutes" in update_data and update_data["estimated_duration_minutes"] is not None:
        paper.estimated_duration_minutes = update_data["estimated_duration_minutes"]
        updated_fields.append("estimated_duration_minutes")
    
    db.commit()
    
    return {
        "paper_id": paper_id,
        "updated_fields": updated_fields,
        "title": paper.title,
        "direction": paper.direction,
        "difficulty": paper.difficulty.value if paper.difficulty else None,
        "description": paper.description,
        "estimated_duration_minutes": paper.estimated_duration_minutes
    }


def export_paper_to_word_enhanced(
    db: Session,
    paper_id: int,
    teacher_id: int,
    include_answers: bool = False
):
    """增强的试卷导出功能"""
    # 获取试卷详情
    paper_detail = get_paper_enhanced_detail(db, paper_id, teacher_id)
    
    # 这里应该实现Word文档生成逻辑
    # 由于需要依赖python-docx等库，这里先返回基本信息
    # 实际实现时需要安装相关依赖并生成真实的Word文档
    
    import tempfile
    import os
    from datetime import datetime
    
    # 创建临时文件
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"试卷_{paper_detail['title']}_{timestamp}.docx"
    
    # 这里应该生成真实的Word文档
    # 暂时返回文件信息
    return {
        "filename": filename,
        "paper_title": paper_detail['title'],
        "total_score": paper_detail['total_score'],
        "question_count": paper_detail['question_count'],
        "include_answers": include_answers,
        "export_time": datetime.now(),
        "sections": paper_detail['sections']
    } 